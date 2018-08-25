from docx import Document
from docx.shared import Inches


class RecordSet:
    def __iter__(self):
        self.qty = "10"
        self.id = "3393993"
        self.desc = "description"


document = Document()

document.add_heading("Document Title", 0)

p = document.add_paragraph("A plain paragraph having some ")
p.add_run("bold").bold = True
p.add_run(" and some ")
p.add_run("italic.").italic = True

document.add_heading("Heading, level 1", level=1)
document.add_paragraph("Intense quote", style="IntenseQuote")

document.add_paragraph("first item in unordered list", style="ListBullet")
document.add_paragraph("first item in ordered list", style="ListNumber")

# document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Qty"
hdr_cells[1].text = "Id"
hdr_cells[2].text = "Desc"

recordset = {"id": 1, "qty": 2, "desc": "descript"}

row_cells = table.add_row().cells
for i, (key, value) in enumerate(recordset.items()):
    row_cells[i].text = str(value)

document.add_page_break()

document.save("demo_row_up.docx")
